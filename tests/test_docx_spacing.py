"""
Script de test para analizar espaciado de acordes en archivos DOCX
Analiza las primeras líneas y muestra información detallada sobre el espaciado
"""

import sys
import os

# Agregar el directorio raíz al path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document as DocxDocument
from core.font_converter import FontConverter
from core.file_processor import FileProcessor
from collections import defaultdict

def analyze_all_fonts(file_path: str):
    """
    Analizar todas las tipografías presentes en el documento
    
    Returns:
        Dict con estadísticas de cada tipografía {font_key: {count, percentage, sample}}
    """
    try:
        doc = DocxDocument(file_path)
        font_usage = defaultdict(lambda: {'count': 0, 'chars': 0, 'sample': ''})
        total_runs = 0
        total_chars = 0
        
        # Analizar cada run en cada párrafo
        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text.strip()
                if not text:
                    continue
                
                total_runs += 1
                total_chars += len(text)
                
                # Obtener nombre y tamaño de fuente
                font_name = run.font.name or 'Arial'
                font_size = int(run.font.size.pt) if run.font.size else 11
                
                font_key = f"{font_name}|{font_size}"
                font_usage[font_key]['count'] += 1
                font_usage[font_key]['chars'] += len(text)
                
                # Guardar muestra de texto (primera aparición)
                if not font_usage[font_key]['sample']:
                    font_usage[font_key]['sample'] = text
        
        # Calcular porcentajes y ordenar por uso
        font_stats = {}
        for font_key, data in font_usage.items():
            percentage = (data['chars'] / total_chars * 100) if total_chars > 0 else 0
            font_stats[font_key] = {
                'count': data['count'],
                'chars': data['chars'],
                'percentage': percentage,
                'sample': data['sample']
            }
        
        # Ordenar por cantidad de caracteres (más usado primero)
        font_stats = dict(sorted(
            font_stats.items(), 
            key=lambda x: x[1]['chars'], 
            reverse=True
        ))
        
        return font_stats
        
    except Exception as e:
        print(f"❌ Error analizando tipografías: {e}")
        return {}


def analyze_docx_spacing(file_path: str, num_lines: int = 10):
    """
    Analizar el espaciado de las primeras líneas de un archivo DOCX
    
    Args:
        file_path: Ruta al archivo .docx
        num_lines: Número de líneas a analizar
    """
    print("="*80)
    print(f"📄 ANÁLISIS DE ESPACIADO: {os.path.basename(file_path)}")
    print("="*80)
    
    # 1. Leer el documento
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        print(f"❌ Error abriendo documento: {e}")
        return
    
    # 2. Analizar TODAS las tipografías del documento
    font_converter = FontConverter()
    
    print(f"\n🔍 ANÁLISIS DE TIPOGRAFÍAS EN EL DOCUMENTO:")
    print("-"*80)
    
    font_stats = analyze_all_fonts(file_path)
    
    if not font_stats:
        print("❌ No se encontraron tipografías en el documento")
        return
    
    # Mostrar todas las tipografías encontradas
    print("\nTipografías encontradas:")
    for i, (font_key, info) in enumerate(font_stats.items(), 1):
        font_name, font_size = font_key.split('|')
        print(f"  {i}. {font_name} {font_size}pt - {info['count']} usos ({info['percentage']:.1f}%)")
        if info['sample']:
            sample_preview = info['sample'][:50] + "..." if len(info['sample']) > 50 else info['sample']
            print(f"     Ejemplo: \"{sample_preview}\"")
    
    # Permitir elegir tipografía
    print(f"\n💬 ¿Qué tipografía usar para el análisis?")
    choice = input(f"   Ingresa número (1-{len(font_stats)}) o Enter para usar la más común: ").strip()
    
    if choice and choice.isdigit():
        choice_idx = int(choice) - 1
        if 0 <= choice_idx < len(font_stats):
            selected_key = list(font_stats.keys())[choice_idx]
        else:
            print("   ⚠️  Opción inválida, usando la más común")
            selected_key = list(font_stats.keys())[0]
    else:
        print("   ℹ️  Usando la tipografía más común")
        selected_key = list(font_stats.keys())[0]
    
    font_name, font_size = selected_key.split('|')
    detected_font = {
        'name': font_name,
        'size': int(font_size),
        'confidence': font_stats[selected_key]['percentage'] / 100
    }
    
    print(f"\n✅ Tipografía seleccionada: {font_name} {font_size}pt")
    
    # 3. Extraer primeras líneas
    paragraphs = [p.text for p in doc.paragraphs if p.text is not None]
    lines = "\n".join(paragraphs).split('\n')[:num_lines]
    
    print(f"\n📝 TEXTO ORIGINAL (primeras {num_lines} líneas):")
    print("-"*80)
    for i, line in enumerate(lines, 1):
        # Mostrar espacios visibles
        visible_line = line.replace(' ', '·')
        print(f"{i:2d}: {visible_line}")
        print(f"    Largo: {len(line)} caracteres")
    
    # 4. Analizar cada línea con detalles
    print(f"\n🔬 ANÁLISIS DETALLADO DE ESPACIADO:")
    print("-"*80)
    
    for i, line in enumerate(lines, 1):
        if not line.strip():
            continue
            
        print(f"\nLínea {i}: '{line}'")
        print(f"  Representación visual: '{line.replace(' ', '·')}'")
        
        # Analizar segmentos
        segments = []
        current_word = []
        current_spaces = []
        
        for char in line:
            if char == ' ':
                if current_word:
                    segments.append(('word', ''.join(current_word)))
                    current_word = []
                current_spaces.append(char)
            else:
                if current_spaces:
                    segments.append(('space', len(current_spaces)))
                    current_spaces = []
                current_word.append(char)
        
        if current_word:
            segments.append(('word', ''.join(current_word)))
        if current_spaces:
            segments.append(('space', len(current_spaces)))
        
        # Mostrar segmentos
        print(f"  Segmentos:")
        for seg_type, seg_value in segments:
            if seg_type == 'word':
                print(f"    - Palabra: '{seg_value}'")
            else:
                print(f"    - Espacios: {seg_value}")
    
    # 5. Mostrar métricas de caracteres disponibles
    print(f"\n📊 MÉTRICAS DE CARACTERES DISPONIBLES:")
    print("-"*80)
    
    # Verificar si hay métricas para la tipografía seleccionada
    if font_name in font_converter.metrics_cache:
        print(f"\n✓ Métricas encontradas para {font_name}")
        available_sizes = sorted(font_converter.metrics_cache[font_name].keys())
        print(f"  Tamaños disponibles: {available_sizes}")
        
        if int(font_size) in font_converter.metrics_cache[font_name]:
            print(f"  ✓ Tamaño {font_size}pt tiene métricas exactas")
            metrics = font_converter.metrics_cache[font_name][int(font_size)]
            print(f"  Caracteres con métricas: {list(metrics.keys())}")
        else:
            print(f"  ⚠️  Tamaño {font_size}pt NO tiene métricas exactas")
            print(f"  Se usará interpolación o el tamaño más cercano")
    else:
        print(f"\n⚠️  NO hay métricas para {font_name}")
        print(f"  Se usará el valor por defecto (0.6)")
    
    # Mostrar ejemplos de anchos de caracteres comunes
    print(f"\n  Ejemplos de anchos relativos (1.0 = Courier New):")
    test_chars = ['i', 'l', 'm', 'w', ' ', 'a', 'e', 'o', 'A', 'E']
    for char in test_chars:
        width = font_converter.get_char_width(char, font_name, int(font_size))
        bar = '█' * int(width * 20)  # Barra visual
        print(f"    '{char}': {width:.3f} {bar}")
    
    # 6. Calcular anchos monoespaciados con la tipografía seleccionada
    print(f"\n📐 CONVERSIÓN A MONOESPACIADO ({font_name} {font_size}pt):")
    print("-"*80)
    
    font_info = {
        'name': font_name,
        'size': int(font_size)
    }
    
    for i, line in enumerate(lines[:5], 1):  # Solo primeras 5 para no saturar
        if not line.strip():
            continue
            
        print(f"\nLínea {i}:")
        print(f"  Original ({len(line)} chars): '{line.replace(' ', '·')}'")
        
        # Calcular ancho total en unidades monospace
        total_width = font_converter.calculate_monospace_width(
            line, 
            font_info['name'], 
            font_info['size']
        )
        print(f"  Ancho monospace: {total_width:.2f} caracteres")
        
        # Convertir línea
        converted = font_converter._convert_line(
            line,
            font_info['name'],
            font_info['size']
        )
        print(f"  Convertido ({len(converted)} chars): '{converted.replace(' ', '·')}'")
        
        # Comparar diferencias
        diff = len(converted) - len(line)
        print(f"  Diferencia: {diff:+d} caracteres")
    
    # 7. Test de alineación de acordes
    print(f"\n🎵 TEST DE ALINEACIÓN DE ACORDES:")
    print("-"*80)
    
    # Buscar par de líneas (acordes + letra)
    file_processor = FileProcessor()
    for i in range(len(lines) - 1):
        line1 = lines[i]
        line2 = lines[i + 1]
        
        if file_processor._is_chord_line(line1) and not file_processor._is_chord_line(line2):
            print(f"\n✓ Par encontrado (líneas {i+1} y {i+2}):")
            print(f"\n  Acordes originales:")
            print(f"    '{line1.replace(' ', '·')}'")
            print(f"  Letra original:")
            print(f"    '{line2.replace(' ', '·')}'")
            
            # Aplicar alineación
            chord_aligned, lyric_padded = file_processor.align_chord_over_lyric(
                line1, line2
            )
            
            print(f"\n  Acordes alineados:")
            print(f"    '{chord_aligned.replace(' ', '·')}'")
            print(f"  Letra ajustada:")
            print(f"    '{lyric_padded.replace(' ', '·')}'")
            
            print(f"\n  Cambios:")
            print(f"    Acordes: {len(line1)} → {len(chord_aligned)} chars ({len(chord_aligned)-len(line1):+d})")
            print(f"    Letra:   {len(line2)} → {len(lyric_padded)} chars ({len(lyric_padded)-len(line2):+d})")
            
            break  # Solo analizar el primer par
    
    print("\n" + "="*80)
    print("✅ Análisis completado")
    print("="*80)


def main():
    """Función principal"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Analizar espaciado de acordes en archivos DOCX'
    )
    parser.add_argument(
        'file',
        help='Ruta al archivo .docx a analizar'
    )
    parser.add_argument(
        '-n', '--lines',
        type=int,
        default=10,
        help='Número de líneas a analizar (default: 10)'
    )
    
    args = parser.parse_args()
    
    if not os.path.exists(args.file):
        print(f"❌ Error: El archivo '{args.file}' no existe")
        return
    
    if not args.file.lower().endswith('.docx'):
        print(f"⚠️  Advertencia: El archivo no tiene extensión .docx")
    
    analyze_docx_spacing(args.file, args.lines)


if __name__ == "__main__":
    # Si se ejecuta sin argumentos, usar archivo de ejemplo
    if len(sys.argv) == 1:
        print("💡 Uso: python test_docx_spacing.py <archivo.docx> [-n NUM_LINEAS]")
        print("\nEjemplo:")
        print('  python test_docx_spacing.py "Al altar del Señor.docx"')
        print('  python test_docx_spacing.py "cancion.docx" -n 15')
        sys.exit(1)
    
    main()
